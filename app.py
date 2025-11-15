# app.py — Final merged Streamlit app (Option C: text + table; exact Tkinter formulas)
import streamlit as st
import networkx as nx
import matplotlib.pyplot as plt
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
from matplotlib.patches import FancyArrow, Rectangle
import numpy as np
import math
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os
import pandas as pd

# ---------------- Page setup ----------------
st.set_page_config(page_title="Network Topology Simulator", layout="wide")

# ---------------- Assets loader ----------------
ASSET_DIR = "assets"

def load_asset_image(filename):
    path = os.path.join(ASSET_DIR, filename)
    try:
        return Image.open(path).convert("RGBA")
    except Exception:
        return None

computer_img = load_asset_image("computer.png")
prof_img = load_asset_image("professor.png")
abijith_img = load_asset_image("abijith.jpeg")
dharmayu_img = load_asset_image("dharmyu.jpeg")

# ---------------- Helpers (kept same as original logic) ----------------
def hierarchy_pos(G, root, width=2., vert_gap=0.4, vert_loc=0, xcenter=0.5, pos=None, parent=None):
    if pos is None:
        pos = {root: (xcenter, vert_loc)}
    else:
        pos[root] = (xcenter, vert_loc)
    children = list(G.neighbors(root))
    if parent is not None and parent in children:
        children.remove(parent)
    if len(children) != 0:
        dx = width / len(children)
        nextx = xcenter - width / 2 - dx / 2
        for child in children:
            nextx += dx
            pos = hierarchy_pos(G, child, width=dx, vert_gap=vert_gap,
                                vert_loc=vert_loc - vert_gap, xcenter=nextx, pos=pos, parent=root)
    return pos

def draw_double_edge_with_arrows(ax, x1, y1, x2, y2, color="#444"):
    dx = x2 - x1
    dy = y2 - y1
    L = math.hypot(dx, dy)
    if L == 0:
        return
    nxp = -dy / L
    nyp = dx / L
    offset = 0.06
    ax.plot([x1 + nxp*offset, x2 + nxp*offset], [y1 + nyp*offset, y2 + nyp*offset], color=color, linewidth=2, zorder=1)
    ax.plot([x1 - nxp*offset, x2 - nxp*offset], [y1 - nyp*offset, y2 - nyp*offset], color=color, linewidth=2, zorder=1)
    mid1x = (x1 + x2)/2 + nxp*offset
    mid1y = (y1 + y2)/2 + nyp*offset
    arr1 = FancyArrow(mid1x, mid1y, dx*0.0001, dy*0.0001, width=0.008, head_width=0.05, head_length=0.05, color='black', length_includes_head=True, zorder=3)
    ax.add_patch(arr1)
    mid2x = (x1 + x2)/2 - nxp*offset
    mid2y = (y1 + y2)/2 - nyp*offset
    arr2 = FancyArrow(mid2x, mid2y, -dx*0.0001, -dy*0.0001, width=0.008, head_width=0.05, head_length=0.05, color='black', length_includes_head=True, zorder=3)
    ax.add_patch(arr2)

def build_step_by_step(n, num_edges, port_cost, cable_len, cable_cost):
    # This mirrors your original Tkinter build_step_by_step exactly
    lines = []
    lines.append("Step-by-step cost calculation:")
    lines.append(f"1) Number of nodes = {n}")
    lines.append(f"2) Number of connections (edges) = {num_edges}")
    lines.append(f"3) Cost per port = {port_cost}")
    lines.append(f"4) Total port cost = connections * 2 * cost_per_port")
    total_port_cost = num_edges * 2 * port_cost
    lines.append(f"   = {num_edges} * 2 * {port_cost} = {total_port_cost:.2f}")
    lines.append(f"5) Cable length per connection = {cable_len} m")
    lines.append(f"6) Cost per unit cable = {cable_cost} per m")
    total_cable_cost = num_edges * cable_len * cable_cost
    lines.append(f"7) Total cable cost = connections * cable_len * cable_cost")
    lines.append(f"   = {num_edges} * {cable_len} * {cable_cost} = {total_cable_cost:.2f}")
    total_cost = total_port_cost + total_cable_cost
    lines.append(f"8) Total cost = total_port_cost + total_cable_cost = {total_port_cost:.2f} + {total_cable_cost:.2f} = {total_cost:.2f}")
    return "\n".join(lines), total_port_cost, total_cable_cost, total_cost

# ---------------- Sidebar controls ----------------
st.sidebar.title("Controls")
with st.sidebar.form("inputs"):
    nodes = st.number_input("Number of Nodes", min_value=1, value=6, step=1)
    topology = st.selectbox("Select Topology", ["Bus", "Star", "Ring", "Mesh", "Tree"])
    ring_variant = st.selectbox("Ring Variant", ["Singly (Bidirectional)", "Singly (Unidirectional)", "Doubly (Unidirectional)"])
    st.sidebar.markdown("---")
    port_cost = st.number_input("Cost per Port (₹)", value=100.0, min_value=0.0, step=1.0, format="%.2f")
    cable_len = st.number_input("Cable Length per Connection (m)", value=10.0, min_value=0.0, step=0.5, format="%.2f")
    cable_cost = st.number_input("Cost per Unit Cable (₹/m)", value=50.0, min_value=0.0, step=1.0, format="%.2f")
    st.sidebar.markdown("---")
    generate = st.form_submit_button("Generate Topology")

# ---------------- Layout (wider simulator) ----------------
col1, col2 = st.columns([3, 1])

# ---------------- Main simulator area ----------------
with col1:
    st.header("Network Topology Simulator 🛠")

    if not generate:
        st.info("Adjust parameters and click **Generate Topology**.")
        st.stop()

    # Build graph exactly as original
    if topology == "Ring" and ring_variant in ["Singly (Unidirectional)", "Doubly (Unidirectional)"]:
        G = nx.DiGraph()
    else:
        G = nx.Graph()
    G.add_nodes_from(range(1, int(nodes) + 1))

    if topology == "Bus":
        for i in range(1, int(nodes)):
            G.add_edge(i, i+1)
    elif topology == "Star":
        for i in range(2, int(nodes) + 1):
            G.add_edge(1, i)
    elif topology == "Ring":
        for i in range(1, int(nodes)):
            G.add_edge(i, i+1)
        G.add_edge(int(nodes), 1)
        if ring_variant == "Doubly (Unidirectional)":
            for i in range(1, int(nodes)):
                G.add_edge(i+1, i)
            G.add_edge(1, int(nodes))
    elif topology == "Mesh":
        for i in range(1, int(nodes) + 1):
            for j in range(i+1, int(nodes) + 1):
                G.add_edge(i, j)
    elif topology == "Tree":
        k = math.log2(int(nodes) + 1)
        if not k.is_integer():
            st.error("Number of nodes for Tree must be 2^k - 1 (e.g., 3, 7, 15...) for a perfect binary tree.")
            st.stop()
        for i in range(2, int(nodes) + 1):
            G.add_edge(i // 2, i)

    # cost calculation using original formulas (num_edges based)
    num_edges = G.number_of_edges()
    step_text, total_port_cost, total_cable_cost, total_cost = build_step_by_step(int(nodes), num_edges, float(port_cost), float(cable_len), float(cable_cost))

    # --- ORIGINAL SUMMARY TEXT (Tkinter style) ---
    summary = (
        f"Topology: {topology} {f'({ring_variant})' if topology=='Ring' else ''}\n"
        f"Nodes: {int(nodes)}\n"
        f"Connections: {num_edges}\n"
        f"Total Cost: ₹{total_cost:.2f}\n\n"
        f"{step_text}\n"
    )

    st.subheader("Results (original format)")
    st.text_area("Summary & Step-by-step", value=summary, height=260)

    # --- Additional Streamlit table (Option C) ---
    st.subheader("Cost Breakdown (table)")
    table_data = {
        "Topology": [topology],
        "Ring Variant": [ring_variant if topology=="Ring" else "N/A"],
        "Nodes": [int(nodes)],
        "Connections": [num_edges],
        "Cost per Port (₹)": [float(port_cost)],
        "Total Port Cost (₹)": [total_port_cost],
        "Cable length per connection (m)": [float(cable_len)],
        "Cost per Unit Cable (₹/m)": [float(cable_cost)],
        "Total Cable Cost (₹)": [total_cable_cost],
        "Total Cost (₹)": [total_cost]
    }
    df = pd.DataFrame(table_data)
    st.table(df)

    # --- Prepare positions for drawing ---
    if topology == "Bus":
        node_list = list(G.nodes)
        ncount = len(node_list)
        xs = np.arange(ncount)
        ys = np.array([0.6 if i % 2 == 0 else -0.6 for i in range(ncount)])
        pos = {node: (float(x), float(y)) for node, x, y in zip(node_list, xs, ys)}
    elif topology == "Ring":
        ncount = len(G.nodes)
        theta = np.linspace(0, 2 * np.pi, ncount, endpoint=False)
        pos = {node: (np.cos(th), np.sin(th)) for node, th in zip(G.nodes, theta)}
    elif topology == "Tree":
        pos = hierarchy_pos(G, 1)
    else:
        pos = nx.spring_layout(G, seed=42)

    # --- Draw graph to matplotlib figure ---
    fig, ax = plt.subplots(figsize=(9, 7))
    ax.set_axis_off()

    if topology != "Bus":
        edges_drawn = set()
        for u, v in G.edges():
            x1, y1 = pos[u]
            x2, y2 = pos[v]
            if topology == "Ring" and ring_variant == "Doubly (Unidirectional)":
                pair = tuple(sorted((u, v)))
                if pair not in edges_drawn:
                    draw_double_edge_with_arrows(ax, x1, y1, x2, y2, color="#444")
                    edges_drawn.add(pair)
            else:
                ax.plot([x1, x2], [y1, y2], color="#444", linewidth=2, zorder=1)
                if topology == "Ring" and ring_variant == "Singly (Unidirectional)":
                    dx = x2 - x1; dy = y2 - y1
                    midpoint = (x1 + 0.5 * dx, y1 + 0.5 * dy)
                    arr = FancyArrow(midpoint[0], midpoint[1], dx * 0.0001, dy * 0.0001,
                                     width=0.008, head_width=0.05, head_length=0.05,
                                     color='black', length_includes_head=True, zorder=3)
                    ax.add_patch(arr)
    else:
        left, right = -0.5, len(pos) - 0.5
        ax.plot([left, right], [0, 0], color='black', linewidth=4, zorder=0)
        for tx in [left, right]:
            term_rect = Rectangle((tx - 0.09, -0.11), 0.18, 0.22, linewidth=0, edgecolor=None, facecolor='black', zorder=2)
            ax.add_patch(term_rect)
        for node, (x, y) in pos.items():
            ax.plot([x, x], [y - 0.08 if y > 0 else y + 0.08, 0], color='black', linewidth=2, zorder=1)
            ax.plot([x - 0.06, x + 0.06], [0, 0], color='black', linewidth=2, zorder=1)

    # Draw nodes as images (if available) or fallback to circles
    if computer_img is not None:
        comp_arr = np.array(computer_img)
        node_count = len(pos)
        zoom = 0.12 if node_count <= 8 else (0.09 if node_count <= 16 else 0.06)
        for node, (x, y) in pos.items():
            ab = AnnotationBbox(OffsetImage(comp_arr, zoom=zoom), (x, y), frameon=False, zorder=4)
            ax.add_artist(ab)
            ax.text(x, y + 0.02, str(node),
                    fontsize=12, fontweight='bold',
                    ha='center', va='center',
                    color='white', zorder=6)
    else:
        nx.draw_networkx_nodes(G, pos, ax=ax, node_size=1100, node_color="#b7e2f9")
        nx.draw_networkx_labels(G, pos, ax=ax, font_size=13, font_weight='bold')

    # Tidy up axes
    xs = [p[0] for p in pos.values()] if pos else []
    ys = [p[1] for p in pos.values()] if pos else []
    if xs and ys:
        xmin, xmax = min(xs) - 0.9, max(xs) + 0.9
        ymin, ymax = min(ys) - 0.9, max(ys) + 0.9
        ax.set_xlim(xmin, xmax)
        ax.set_ylim(ymin, ymax)

    ax.set_title(f"{topology}" + (f" ({ring_variant})" if topology == "Ring" else ""), fontsize=15, weight='bold')
    st.subheader("Topology Diagram")
    st.pyplot(fig)

    # ---------------- DOCX report generator (restored behavior) ----------------
    def create_docx_bytes():
        doc = Document()
        # Title
        title = doc.add_heading('Network Topology Simulator - Detailed Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            title.runs[0].font.color.rgb = RGBColor(0, 122, 204)
        except Exception:
            pass

        doc.add_paragraph()
        doc.add_heading('Input Parameters', 1)
        input_table = doc.add_table(rows=6, cols=2)
        input_table.style = 'Light Grid Accent 1'
        input_data = [
            ('Topology Type', f"{topology} {('('+ring_variant+')') if topology=='Ring' else ''}"),
            ('Number of Nodes', str(nodes)),
            ('Cost per Port (₹)', f"{port_cost:.2f}"),
            ('Cable Length per Connection (m)', f"{cable_len:.2f}"),
            ('Cost per Unit Cable (₹/m)', f"{cable_cost:.2f}"),
            ('Total Connections', str(num_edges))
        ]
        for i, (param, value) in enumerate(input_data):
            input_table.rows[i].cells[0].text = param
            input_table.rows[i].cells[1].text = value
            input_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()
        doc.add_heading('Cost Analysis Results', 1)
        result_para = doc.add_paragraph()
        result_para.add_run(summary).font.name = 'Consolas'
        doc.add_paragraph()
        doc.add_heading('Network Topology Diagram', 1)
        img_stream = BytesIO()
        fig.savefig(img_stream, format='PNG', bbox_inches='tight', dpi=150)
        img_stream.seek(0)
        try:
            doc.add_picture(img_stream, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Error embedding graph image: {e}]")
        doc.add_paragraph()
        doc.add_heading('Network Connections', 1)
        edges_para = doc.add_paragraph()
        for u, v in G.edges():
            edges_para.add_run(f"• Node {u} ↔ Node {v}\n")
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        try:
            footer_para.add_run('\n' + '─' * 70 + '\n').font.color.rgb = RGBColor(128, 128, 128)
        except Exception:
            pass
        footer_run = footer_para.add_run('Generated by Network Topology Simulator\n')
        footer_run.font.size = Pt(9)
        try:
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
        except Exception:
            pass
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    st.download_button("📥 Download Detailed Report (.docx)", data=create_docx_bytes(),
                       file_name="network_topology_report.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---------------- Right panel — collapsible about/developers/guide ----------------
with col2:
    with st.expander("📌 About This Project"):
        st.write("Network Topology Simulator — visualizes topologies and computes costs (ports + cables).")
        st.write("Reference video: https://www.youtube.com/watch?v=zbqrNg4C98U")

    with st.expander("👨‍💻 Developers"):
        c1, c2 = st.columns(2)
        with c1:
            if abijith_img:
                st.image(abijith_img, width=140, caption="Abijith Thennarasu")
            else:
                st.write("Abijith Thennarasu")
        with c2:
            if dharmayu_img:
                st.image(dharmayu_img, width=140, caption="Dharmayu Jadwani")
            else:
                st.write("Dharmayu Jadwani")

    with st.expander("🎓 Guide / Faculty"):
        if prof_img:
            st.image(prof_img, width=180, caption="Dr. Swaminathan Annadurai")
        else:
            st.write("Guided by: Dr. Swaminathan Annadurai")
